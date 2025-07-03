import streamlit as st
import json
import os
import re
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from io import BytesIO

# ---------------- CONFIGURAÇÕES ----------------
st.set_page_config(page_title="📄 Gerador de Peças Jurídicas", layout="centered")

API_KEY = "AIzaSyBrRaTvJ6wo2-PQUxOwRMQzr1S8KucT79A"
genai.configure(api_key=API_KEY)
MODEL = "gemini-1.5-flash-latest"

PROMPTS_DIR = "Prompts/Pecas"

# ---------------- FUNÇÕES AUXILIARES ----------------
def carregar_areas():
    return [f.replace(".json", "") for f in os.listdir(PROMPTS_DIR) if f.endswith(".json")]

def carregar_json_area(area):
    with open(f"{PROMPTS_DIR}/{area}.json", "r", encoding="utf-8") as f:
        return json.load(f)

def extrair_conteudo(arquivo):
    if arquivo.name.endswith(".pdf"):
        reader = PdfReader(arquivo)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text() + "\n"
        return texto
    elif arquivo.name.endswith(".docx"):
        doc = Document(arquivo)
        texto = ""
        for para in doc.paragraphs:
            texto += para.text + "\n"
        return texto
    elif arquivo.name.endswith(".txt"):
        return arquivo.read().decode("utf-8", errors="ignore")
    else:
        return ""

# ---------------- ESTADO ----------------
if "resposta" not in st.session_state:
    st.session_state.resposta = None

if "exportar" not in st.session_state:
    st.session_state.exportar = False

# ---------------- INÍCIO ----------------
st.title("📄 Gerador de Peças Jurídicas - Campos Extras Controlados")

areas = carregar_areas()
area_escolhida = st.selectbox("Escolha a área do Direito:", ["Selecione..."] + areas)

if area_escolhida != "Selecione...":
    arvore_area = carregar_json_area(area_escolhida)

    selections = {}
    nivel_atual = arvore_area
    nivel = 0

    while True:
        if isinstance(nivel_atual, dict) and "campos" not in nivel_atual:
            opcoes = list(nivel_atual.keys())
            escolha = st.selectbox(f"Nível {nivel+1}:", ["Selecione..."] + opcoes, key=f"nivel_{nivel}")
            if escolha != "Selecione...":
                selections[f"nivel_{nivel}"] = escolha
                nivel_atual = nivel_atual[escolha]
                nivel += 1
            else:
                break
        else:
            break

    if isinstance(nivel_atual, dict) and "campos" in nivel_atual:
        st.header(f"📌 Peça: {selections[f'nivel_{nivel-1}']}")

        entradas = {}

        if nivel_atual.get("permite_upload", False):
            arquivo = st.file_uploader("📎 Anexe documento de exemplo (opcional)", type=["pdf", "docx", "txt"])
        else:
            arquivo = None

        # Campos fixos
        for campo in nivel_atual["campos"]:
            valor = st.text_input(campo["label"]) if campo["tipo"] == "text" else st.text_area(campo["label"])
            if campo.get("regex") and valor:
                if not re.fullmatch(campo["regex"], valor):
                    st.warning(campo.get("regex_msg", f"Valor inválido para {campo['label']}"))
            entradas[campo["nome"]] = valor

        # Grupos de campos extras conforme JSON
        grupos_extras_data = []
        for grupo in nivel_atual.get("grupos_extras", []):
            st.subheader(grupo["label_grupo"])
            num = st.number_input(f"Quantos '{grupo['label_grupo']}'?", min_value=0, step=1, value=0, key=f"num_{grupo['nome_grupo']}")
            grupo_instancias = []
            for i in range(num):
                st.markdown(f"**{grupo['label_grupo']} #{i+1}**")
                instancia = {}
                for campo in grupo["campos"]:
                    k = f"{grupo['nome_grupo']}_{i}_{campo['nome']}"
                    if campo["tipo"] == "text":
                        v = st.text_input(campo["label"], key=k)
                    else:
                        v = st.text_area(campo["label"], key=k)
                    if campo.get("regex") and v:
                        if not re.fullmatch(campo["regex"], v):
                            st.warning(campo.get("regex_msg", f"Valor inválido para {campo['label']}"))
                    instancia[campo["nome"]] = v
                grupo_instancias.append(instancia)
            grupos_extras_data.append((grupo["nome_grupo"], grupo_instancias))

        gerar = st.button("📑 Gerar Peça")
        if gerar:
            prompt_base = nivel_atual["prompt"]
            for nome, valor in entradas.items():
                prompt_base = prompt_base.replace(f"{{{{{nome}}}}}", str(valor))

            for nome_grupo, instancias in grupos_extras_data:
                prompt_base += f"\n\n[{nome_grupo.upper()}]\n"
                for i, instancia in enumerate(instancias, 1):
                    prompt_base += f"- {nome_grupo} #{i}:\n"
                    for campo_nome, campo_valor in instancia.items():
                        prompt_base += f"   - {campo_nome}: {campo_valor}\n"

            if arquivo:
                arquivo_content = extrair_conteudo(arquivo)
                prompt_base += f"\n\n[DOCUMENTO DE EXEMPLO]\n{arquivo_content}"

            model = genai.GenerativeModel(MODEL)
            resposta = model.generate_content(prompt_base).text

            st.session_state.resposta = resposta  # Salva no estado
            st.success("✅ Peça Gerada com Sucesso!")
            st.write(resposta)

        # Mostrar exportar para DOCX apenas se houver resposta gerada
        if st.session_state.resposta:
            st.divider()
            st.subheader("📤 Exportar para DOCX")
            st.session_state.exportar = st.checkbox("Quero exportar para DOCX", value=st.session_state.exportar)

            if st.session_state.exportar:
                cabecalho_img = st.file_uploader("🖼️ Inserir imagem do cabeçalho (opcional)", type=["png", "jpg", "jpeg"], key="cabecalho")
                rodape_img = st.file_uploader("🖼️ Inserir imagem do rodapé (opcional)", type=["png", "jpg", "jpeg"], key="rodape")

                buffer = BytesIO()
                if st.button("📥 Clique aqui para baixar DOCX"):
                    doc = Document()
                    section = doc.sections[0]

                    # Cabeçalho
                    if cabecalho_img:
                        header = section.header
                        header_paragraph = header.paragraphs[0]
                        header_paragraph.add_run().add_picture(cabecalho_img, width=Inches(6.0))

                    # Corpo do texto
                    doc.add_paragraph(st.session_state.resposta)

                    # Rodapé
                    if rodape_img:
                        footer = section.footer
                        footer_paragraph = footer.paragraphs[0]
                        footer_paragraph.add_run().add_picture(rodape_img, width=Inches(6.0))

                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="📥 Clique aqui para baixar DOCX",
                        data=buffer,
                        file_name="peca_gerada.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
